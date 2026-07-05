"""Ярус 2: извлечение таблиц составов через мультимодальную модель (Qwen3-VL, RouterAI).
Рендер страницы → OCR → структурные ячейки → факты с ПРАВИЛЬНОЙ привязкой элемент↔значение.
Решает провал плоского текста на безрамочных таблицах составов.
API-based (ноль локального RAM); дисковый кэш OCR-ответов (идемпотентно, не жечь квоту).
Пре-фильтр: OCR шлём только на страницы с признаком таблицы состава (≥3 хим-токена + числа).
"""
from __future__ import annotations
import os, sys, json, re, base64, hashlib, time
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from src.config import DOCS_META, DOCS_TEXT, ARTIFACTS, LLM_ENABLED, nfc
import requests

CACHE_DIR = ARTIFACTS.parent / "data" / "ocr_cache"
CACHE_DIR.mkdir(parents=True, exist_ok=True)
VISION_FACTS = ARTIFACTS / "vision_facts.jsonl"

# --- химические токены (заголовки столбцов таблиц состава) ---
ELEMENTS = {"S","Se","Te","Ni","Cu","Co","Fe","Pb","Zn","As","Sb","Au","Ag","Pt","Pd","Rh",
            "C","Si","Ca","Mg","Al","Na","Cl","O","Cd","Bi","Sn","Cr","Mn","P","N"}
OXIDE = re.compile(r"^[A-Z][a-z]?\d?O\d?$|^(SiO2|Al2O3|Fe2O3|CaO|MgO|NiO|CuO|Cu2O|Na2O|K2O|SO3|P2O5|TiO2|MnO|Cr2O3|ZnO|PbO)$")
ELEM2CANON = {"Ni":"никель","Cu":"медь","Co":"кобальт","Fe":"железо","Au":"золото","Ag":"серебро",
              "Pt":"платина","Pd":"палладий","Rh":"родий","Se":"селен","Te":"теллур","Pb":"свинец",
              "Zn":"цинк","As":"мышьяк","Sb":"сурьма","S":"сера","C":"углерод","Si":"кремний"}
PHASE_KW = {"штейн":"штейн","файнштейн":"файнштейн","шлак":"шлак","шлам":"шлам","раствор":"раствор",
            "католит":"католит","анолит":"анолит","концентрат":"концентрат","огарок":"огарок",
            "кек":"кек","пыл":"пыль","возгон":"возгон"}

# --- OCR-конфузии: латиница↔кириллица, '1'↔'l', '2'↔'Z' и т.п. в хим-символах ---
# Кириллические гомоглифы → латиница (только для сверки с ELEMENTS/OXIDE)
_CYR2LAT = str.maketrans({
    "А": "A", "В": "B", "С": "C", "Е": "E", "Н": "H", "К": "K", "М": "M",
    "О": "O", "Р": "P", "Т": "T", "Х": "X", "а": "a", "е": "e", "о": "o",
    "с": "c", "р": "p", "у": "y", "х": "x", "Ѕ": "S", "І": "I", "Ғ": "F",
    "г": "r", "Г": "r",
})
# Прямые замены известных OCR-искажений хим-символов.
# ВАЖНО: НЕТ самостоятельных '0'/'О'→'O' — одиночный '0' это ЧИСЛО (ноль),
# а не кислород; иначе столбец '0|0|0|0' даёт ложные факты кислорода (issue #1).
# Замена цифра→O делается только внутри токена длиной ≥3, если после замены
# получается валидный OXIDE (Si02→SiO2, A12O3→Al2O3).
_CONFUSIONS = {
    "AI": "Al", "A1": "Al", "AL": "Al", "A12O3": "Al2O3", "AI2O3": "Al2O3",
    "Fе": "Fe", "Ғе": "Fe", "Ғе2О3": "Fe2O3", "Fe2О3": "Fe2O3",
    "Сг": "Cr", "СГ": "Cr", "Со": "Co", "СО": "Co", "Си": "Cu",
    "Ni": "Ni", "N1": "Ni", "Nl": "Ni", "SiО2": "SiO2", "Si02": "SiO2",
    "SO3": "SO3", "S03": "SO3",
}

def _digit_to_oxide(t: str):
    """Цифра→O ТОЛЬКО внутри токена длиной ≥3, если после замены '0'/'О'→'O'
    получается валидный оксид (Si02→SiO2). Одиночные '0'/'О' не трогаем."""
    if len(t) < 3:
        return None
    cand = t.replace("0", "O").replace("О", "O")
    if cand != t and OXIDE.match(cand):
        return cand
    return None

def _clean_tok(tok: str) -> str:
    """Чистка OCR-токена перед сверкой с ELEMENTS: strip мусора, гомоглыфы, конфузии."""
    t = tok.strip().strip(",;:.").strip(" _").strip("|").strip()
    # хвостовые '-' / '_' от переноса или разделителя
    t = t.rstrip("-_ ")
    if not t:
        return t
    # прямая замена по словарю искажений (до транслитерации)
    if t in _CONFUSIONS:
        return _CONFUSIONS[t]
    # цифра→O внутри оксида (Si02→SiO2), но НЕ для одиночного '0'
    ox = _digit_to_oxide(t)
    if ox:
        return ox
    # транслитерация кириллических гомоглифов → латиница
    tl = t.translate(_CYR2LAT)
    if tl in _CONFUSIONS:
        return _CONFUSIONS[tl]
    ox = _digit_to_oxide(tl)
    if ox:
        return ox
    if tl in ELEMENTS or OXIDE.match(tl):
        return tl
    return t

def _is_elem(tok: str) -> bool:
    t = _clean_tok(tok)
    return t in ELEMENTS or bool(OXIDE.match(t))

# issue #1: value-ячейка, текст которой содержит БУКВЫ (кириллицу/слова),
# — это метка/подпись ('Серия №N','Номер опыта','Расход газа', тире-с-текстом),
# а НЕ число. Разрешены лишь символы «чистого числа»: цифры, знак, диапазон,
# разделитель, компаратор и единичный «хвост» единицы (%, ‰, º/o C и т.п.).
# Любая латинская/кириллическая буква → ячейка НЕ числовая → значение не берём.
_LETTER_RE = re.compile(r"[A-Za-zА-Яа-яЁё]")
# ведущий компаратор value-ячейки (issue #2): <, >, ≤, ≥, ±, а также =/~
# НОРМАЛИЗАЦИЯ: ≤/≥ → <=/>= — грамматика (grammar.py) хранит ASCII-формы,
# единый словарь компараторов для фильтров поиска.
_COMP_LEAD_RE = re.compile(r"^\s*(<=|>=|≤|≥|±|<|>|=|~|≈)")
_COMP_MAP = {"<=": "<=", ">=": ">=", "≤": "<=", "≥": ">=", "±": "±",
             "<": "<", ">": ">", "=": "=", "~": "≈", "≈": "≈"}

def _is_numeric_cell(cell: str) -> bool:
    """True, если текст value-ячейки — «чистое число» (± знак/диапазон/единица/
    компаратор), а НЕ метка/подпись. Наличие любой буквы → False (issue #1)."""
    s = (cell or "").strip()
    if not s:
        return False
    if _LETTER_RE.search(s):
        return False
    return bool(re.search(r"\d", s))

def _cell_comparator(cell: str):
    """Ведущий компаратор value-ячейки (issue #2): '<0,5'→'<', '≥30'→'>=',
    '±2'→'±'. ≤/≥ нормализуются в <=/>= (единый словарь с grammar.py).
    Нет ведущего оператора → None (по умолчанию '=')."""
    m = _COMP_LEAD_RE.match(cell or "")
    if not m:
        return None
    return _COMP_MAP.get(m.group(1))

# --- диапазон ВНУТРИ одной ячейки: «5-15», «24,19–32,57», «1,0÷2,0», «30-44%» ---
# Дефис '-' добавлен в разделители С ГВАРДАМИ, чтобы НЕ сломать фикс «ID-минус»
# раунда-2 (_num: дефис после буквы/цифры = разделитель ID «S-1281», не минус):
#   • ячейка ЦЕЛИКОМ = число-разделитель-число (± хвост %/‰) — анкера ^…$;
#   • букв в ячейке нет вообще (_is_numeric_cell) → «CUJ-6», «ОЭП-1» не пройдут;
#   • не более одного дефиса → «-5-15», «1-2-3» не диапазон;
#   • обе стороны — 4-значные целые 1900..2100 → это годы (дата), не диапазон.
_CELL_RANGE_RE = re.compile(
    r"^\s*(\d+(?:[.,]\d+)?)\s*(–|—|÷|…|\.\.\.|-)\s*(\d+(?:[.,]\d+)?)\s*[%‰]?\s*$")

def _cell_range(cell: str):
    """(lo, hi) если ячейка — диапазон «X-Y»/«X–Y»/«X÷Y»/«X…Y», иначе None.
    «17,7-» (висячий дефис), «-3,5» (минус), «CUJ-6» (ID) → None."""
    s = (cell or "").strip()
    if not _is_numeric_cell(s):        # буквы → ID/метка, не диапазон
        return None
    m = _CELL_RANGE_RE.match(s)
    if not m:
        return None
    if m.group(2) == "-" and s.count("-") > 1:   # >1 дефиса — ID/дата/список
        return None
    a = float(m.group(1).replace(",", "."))
    b = float(m.group(3).replace(",", "."))
    # пара «годов» (оба целые 1900..2100) — период, не диапазон значений
    if ("," not in m.group(1) and "." not in m.group(1)
            and "," not in m.group(3) and "." not in m.group(3)
            and 1900 <= a <= 2100 and 1900 <= b <= 2100):
        return None
    return (min(a, b), max(a, b))

def _num(cell: str):
    """'17,7-' -> 17.7 ; '20,9' -> 20.9 ; '-3,5' -> -3.5 ; мусор -> None.
    Дефис считается знаком минуса ТОЛЬКО если перед ним не стоит буква/цифра —
    иначе это внутренний разделитель ID ('S-1281','R-1511','ОЭП-1'), не минус.
    ПРИМЕЧАНИЕ: сам _num допускает буквы (нужно для детекции ID/реестров в
    _table_is_composition и data_rows). Отбраковка «ячейка-метка» (issue #1)
    делается на СТОРОНЕ извлечения значения — см. _value_num/_emit_series."""
    s = cell.replace(" ", "")
    for m in re.finditer(r"-?\d+(?:[.,]\d+)?", s):
        g = m.group()
        if g.startswith("-"):
            prev = s[m.start() - 1] if m.start() > 0 else ""
            if prev.isalnum():
                # дефис — разделитель внутри ID, а не минус: берём модуль
                g = g[1:]
        return float(g.replace(",", "."))
    return None

def _value_num(cell: str):
    """Число ИЗ VALUE-ЯЧЕЙКИ: как _num, но ячейка-метка/подпись отбраковывается
    (issue #1). Если текст содержит буквы (кириллицу/слова — 'Серия №1',
    'Номер опыта', 'Расход газа', тире-с-текстом) — это НЕ значение → None.
    Компаратор (<,>,≤,≥,±) перед числом буквой не считается."""
    if not _is_numeric_cell(cell):
        return None
    return _num(cell)

def _canon(tok: str) -> str:
    t = _clean_tok(tok)
    return ELEM2CANON.get(t, t)

# слова-агрегаты в col0 транспонированной/обычной таблицы — не пробы
_AGG_RE = re.compile(r"\b(макс|мин|сред|maximum|minimum|max|min|avg|average|сумм|итог|total)",
                     re.IGNORECASE)
def _is_agg_label(cell: str) -> bool:
    return bool(_AGG_RE.search(cell or ""))

# ------------------------------------------------------------ OCR + кэш ---
# --------------------------------------------------- таблица → факты ---
_MAJOR_OXIDES = {"SiO2","Al2O3","CaO","MgO","Fe2O3","Cr2O3","MnO","Na2O","K2O","TiO2","P2O5","SO3"}

def _has_trailing_dash(cell: str) -> bool:
    return cell.strip().rstrip(" _").endswith("-")

def _emit_series(head_clean: str, cells: list, unit: str, phase, doc_id: str, context: str,
                 allow_pair_range: bool = True, unit_conf: float = 0.95):
    """Из списка сырых ячеек одного элемента строим факты.
    Ячейка-диапазон «5-15»/«24,19–32,57» → факт-диапазон (value_low..value_high).
    >2 РАЗНЫХ значения → факт на каждую пробу (comparator='=', точка), НЕ ложный диапазон.
    2 значения с висячим '-' у первого → диапазон (min-max) — ЕСЛИ allow_pair_range.
    Когда хвостовой '-' СИСТЕМНЫЙ по всей таблице (OCR-перенос), allow_pair_range=False
    и 2 значения дают 2 ТОЧЕЧНЫХ факта, а не ложный диапазон (issue #3).
    1 значение → точка. Ячейки-агрегаты (Макс/Мин/Среднее) пропускаются.
    unit_conf — уверенность в единице pct (0.6, если '%' не подтверждён заголовком)."""
    # issue #5: кислород O — балансный элемент состава, не эмитим как факт
    # (обычно «остаток до 100%», плюс частый ложняк от OCR-путаницы 0↔O)
    if head_clean == "O":
        return []
    if _OCR_ARTIFACT.match(head_clean or ""):   # D5O/D9O — OCR-метка размера частиц, не оксид
        return []
    vals, raws, comps, ranges = [], [], [], []
    for cell in cells:
        if _is_agg_label(cell):
            continue
        # ячейка-диапазон («5-15», «24,19–32,57») → сразу диапазонный факт
        rng = _cell_range(cell)
        if rng is not None:
            ranges.append((rng[0], rng[1], cell))
            continue
        # issue #1: value-ячейка с буквами (метка/подпись) — НЕ число, пропускаем
        v = _value_num(cell)
        if v is not None:
            vals.append(v); raws.append(cell)
            comps.append(_cell_comparator(cell))  # issue #2: ведущий <,>,≤,≥,±
    if not vals and not ranges:
        return []
    u = "pct" if head_clean in _MAJOR_OXIDES else unit
    canon = ELEM2CANON.get(head_clean, head_clean)

    def _mk(lo, hi, comp, raw_repr):
        if u == "pct" and hi > 100:
            return None
        return {
            "doc_id": doc_id, "node_type": "Material", "canon": canon,
            "metric": "содержание", "unit_canon": u,
            "value_low": lo, "value_high": hi, "comparator": comp,
            "confidence": unit_conf if u == "pct" else 0.95,
            "phase": phase, "conditions": None,
            "source": "vision_ocr", "quote": f"{context.strip()[:80]} | {head_clean}: {raw_repr}",
        }

    facts = []
    # ячейки-диапазоны — по факту на каждую (value_low..value_high)
    for lo, hi, raw in ranges:
        f = _mk(lo, hi, "=", raw)
        if f: facts.append(f)
    distinct = set(vals)
    if not vals:
        pass
    elif len(vals) == 1:
        # issue #2: компаратор из ячейки (по умолчанию '='), не хардкод
        f = _mk(vals[0], vals[0], comps[0] or "=", raws[0])
        if f: facts.append(f)
    elif len(vals) == 2 and allow_pair_range and _has_trailing_dash(raws[0]):
        # висячий дефис у верхней ячейки (НЕ системный) → это диапазон min-max
        f = _mk(min(vals), max(vals), "=", "–".join(raws))
        if f: facts.append(f)
    elif len(vals) == 2 and not allow_pair_range:
        # системный хвостовой '-' = OCR-перенос → 2 точечных факта, не диапазон
        for v, raw, comp in zip(vals, raws, comps):
            f = _mk(v, v, comp or "=", raw)
            if f: facts.append(f)
    elif len(distinct) <= 2 and len(vals) == 2:
        f = _mk(min(vals), max(vals), "=", "–".join(raws))
        if f: facts.append(f)
    else:
        # >2 РАЗНЫХ значения: НЕ диапазон — факт на каждую пробу (точка)
        for v, raw, comp in zip(vals, raws, comps):
            f = _mk(v, v, comp or "=", raw)
            if f: facts.append(f)
    return facts

def _detect_orientation(grid: dict, nrows: int, ncols: int):
    """Возвращает ('row', header_row) для обычной таблицы (элементы в шапке-строке),
    ('col', header_col) для ТРАНСПОНИРОВАННОЙ (элементы вниз по столбцу),
    или (None, None)."""
    # обычная: строка с ≥3 хим-токенами
    for r in range(nrows):
        hits = sum(1 for c in range(ncols) if _is_elem(grid.get((r, c), "")))
        if hits >= 3:
            row_hits, row_r = hits, r
            break
    else:
        row_hits, row_r = 0, None
    # транспонированная: столбец с ≥3 хим-токенами (обычно col0)
    for c in range(ncols):
        hits = sum(1 for r in range(nrows) if _is_elem(grid.get((r, c), "")))
        if hits >= 3:
            col_hits, col_c = hits, c
            break
    else:
        col_hits, col_c = 0, None
    # если оба сработали — выбираем ориентацию с бОльшим числом хим-токенов
    if col_c is not None and col_hits > row_hits:
        return "col", col_c
    if row_r is not None:
        return "row", row_r
    if col_c is not None:
        return "col", col_c
    return None, None

def table_to_facts(grid: dict, nrows: int, ncols: int, doc_id: str, context: str):
    """grid[(r,c)]=text. Возвращает факты состава с привязкой элемент↔значение.
    Поддержаны обычные (элементы в строке-шапке) и транспонированные
    (элементы вниз по столбцу; остальные столбцы — пробы) таблицы."""
    orient, idx = _detect_orientation(grid, nrows, ncols)
    if orient is None:
        return []

    if orient == "row":
        header_row = idx
        title_cells = [grid.get((r, c), "") for r in range(header_row + 1) for c in range(ncols)]
    else:
        header_col = idx
        # заголовок таблицы для транспонированной: ячейки левее хим-столбца + строка 0
        title_cells = [grid.get((r, c), "") for r in range(nrows) for c in range(header_col + 1)]
        title_cells += [grid.get((0, c), "") for c in range(ncols)]
    tbl_title = " ".join(title_cells).lower()
    phase = next((v for k, v in PHASE_KW.items() if k in tbl_title), None)
    unit = "pct"
    if "г/л" in tbl_title or "г/дм" in tbl_title: unit = "mg_L"
    elif "г/т" in tbl_title: unit = "g_t"
    # unit_canon=pct — дефолтное допущение; если '%' НЕТ ни в заголовке таблицы,
    # ни в шапке столбцов (title_cells включает шапку), уверенность в единице 0.6;
    # смешанный заголовок «% и ppm» — тоже 0.6 (единица неоднозначна)
    has_pct_hdr = "%" in tbl_title
    has_ppm_hdr = "ppm" in tbl_title
    pct_conf = 0.95 if (has_pct_hdr and not has_ppm_hdr) else 0.6

    facts = []
    if orient == "row":
        # пробы = строки данных ниже шапки; элементы = столбцы-шапки
        data_rows = [r for r in range(header_row + 1, nrows)
                     if any(_num(grid.get((r, c), "")) is not None for c in range(ncols))]
        elem_cols = [c for c in range(ncols) if _is_elem(grid.get((header_row, c), ""))]
        # issue #3: в 2-строчной таблице системный хвостовой '-' у ПЕРВОЙ строки
        # (у большинства хим-столбцов) = OCR-перенос, а не диапазон → точки
        allow_range = True
        if len(data_rows) == 2 and elem_cols:
            r0 = data_rows[0]
            dashed = sum(1 for c in elem_cols if _has_trailing_dash(grid.get((r0, c), "")))
            if dashed > len(elem_cols) / 2:
                allow_range = False
        for c in elem_cols:
            head = _clean_tok(grid.get((header_row, c), ""))
            cells = [grid.get((r, c), "") for r in data_rows]
            facts += _emit_series(head, cells, unit, phase, doc_id, context,
                                  allow_pair_range=allow_range, unit_conf=pct_conf)
    else:
        # ТРАНСПОНИРОВАННАЯ: элементы = строки в header_col; пробы = остальные столбцы
        elem_rows = [r for r in range(nrows) if _is_elem(grid.get((r, header_col), ""))]
        val_cols = [c for c in range(ncols) if c != header_col
                    and any(_num(grid.get((r, c), "")) is not None for r in elem_rows)]
        for r in elem_rows:
            head = _clean_tok(grid.get((r, header_col), ""))
            cells = [grid.get((r, c), "") for c in val_cols]
            facts += _emit_series(head, cells, unit, phase, doc_id, context,
                                  unit_conf=pct_conf)
    return facts

# --------------------------------------------------- пре-фильтр + прогон ---
_COMP_HINT = re.compile(r"\b(SiO2|Al2O3|Fe2O3|Ni|Cu|Co|Se|Te|штейн|шлак|шлам|файнштейн)\b")

def page_worth_ocr(page_text: str) -> bool:
    """Дёшево: ≥3 хим-намёка и есть числа с запятой (состав)."""
    hints = len(set(_COMP_HINT.findall(page_text)))
    return hints >= 3 and bool(re.search(r"\d+,\d", page_text))

# ──────────────────────────────────────────────────────────────────────────
# VL-дорожка: OCR таблиц состава для загружаемых PDF через мультимодальную модель.
# Для свежих PDF таблицы разбирает Qwen3-VL (RouterAI).
# VL даёт СТРУКТУРУ (элемент/фаза), ЧИСЛА валидируются теми же детерминированными
# хелперами (_value_num/_cell_range/_cell_comparator) — детерминизм цифр сохранён.
# Существующий корпус не трогаем: он уже в OCR-кэше (source='vision_ocr').
# ──────────────────────────────────────────────────────────────────────────
_CANON_SET = set(ELEM2CANON.values())
_OCR_ARTIFACT = re.compile(r"^D[0-9]O$")   # D5O/D9O — OCR-метки размера частиц (D50/D90), не оксид


def _is_ocr_artifact(head: str) -> bool:
    return bool(_OCR_ARTIFACT.match((head or "").strip()))

_VL_PROMPT = (
    "Ты извлекаешь ТАБЛИЦЫ ХИМИЧЕСКОГО СОСТАВА со скана страницы (RU/EN).\n"
    "Верни СТРОГО JSON без пояснений и markdown:\n"
    '{"tables":[{"phase":"<фаза материала: штейн/шлак/раствор/католит/концентрат/'
    'файнштейн/огарок/кек или null>","rows":[{"element":"<символ или название: Ni, '
    'никель, SiO2>","value":"<ДОСЛОВНО как в ячейке: 4,2 или 5-15 или <0,5>",'
    '"unit":"<% или г/т или мг/л или ppm или null>"}]}]}\n'
    "Правила: значение копируй ДОСЛОВНО из ячейки (запятая-разделитель, диапазон "
    "через дефис, ведущие <>≤≥). НЕ вычисляй и НЕ округляй. Нет таблиц состава — "
    '{"tables":[]}.'
)


def _vl_ocr(png_bytes: bytes):
    """PNG страницы → распарсенный JSON таблиц состава через VL-модель (RouterAI).
    Кэш и ретраи — через llm.chat (ключ включает картинку → идемпотентно на изображение).
    None при выключенном VL / ошибке (вызывающий деградирует)."""
    from src.config import VL_MODEL, VL_ENABLED
    if not VL_ENABLED:
        return None
    from src import llm
    b64 = base64.b64encode(png_bytes).decode()
    messages = [{"role": "user", "content": [
        {"type": "text", "text": _VL_PROMPT},
        {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b64}"}}]}]
    try:
        raw = llm.chat(messages, model=VL_MODEL, temperature=0,
                       max_tokens=2000, timeout=90, max_attempts=2)
        return llm._extract_json(raw)
    except Exception:  # noqa: BLE001 — VL недоступен/битый ответ → без таблиц
        return None


def _vl_canon(elem: str):
    """Название/символ элемента от VL → (canon, head-для-цитаты). None — пропустить."""
    e = (elem or "").strip()
    if not e:
        return None, None
    if e in ELEM2CANON:                      # символ 'Ni'
        return ELEM2CANON[e], e
    if OXIDE.match(e):                        # оксид 'SiO2' — материал как есть
        return e, e
    low = e.lower().replace("ё", "е")
    if low in _CANON_SET:                     # русское название 'никель'
        return low, e
    return low, e                             # прочий материал (lowercase)


def _vl_unit(u: str, head: str):
    """Единица от VL → канон. Оксиды/проценты → pct; г/т, мг/л, ppm распознаём."""
    s = (u or "").strip().lower()
    if head in _MAJOR_OXIDES:
        return "pct"
    if "%" in s or "мас" in s or s in ("wt", "wt%", "об.%", "pct"):
        return "pct"
    if "г/т" in s or "g/t" in s or "г/тонн" in s:
        return "g_t"
    if "мг/л" in s or "mg/l" in s or "мг/дм" in s:
        return "mg_L"
    if "ppm" in s:
        return "ppm"
    return "pct"     # таблицы состава по умолчанию — проценты


def _vl_to_facts(parsed, doc_id: str, context: str):
    """Распарсенный VL-JSON → факты (та же схема, что vision_ocr; source='vision_vl').
    Числа парсятся детерминированно; O-баланс и pct>100 отбрасываются как в OCR-пути."""
    facts = []
    if not isinstance(parsed, dict):
        return facts
    ctx = (context or "").strip()[:80]
    for tbl in (parsed.get("tables") or []):
        phase = None
        pv = (tbl.get("phase") or "").strip().lower()
        for kw, canon_ph in PHASE_KW.items():
            if kw in pv:
                phase = canon_ph
                break
        for row in (tbl.get("rows") or []):
            canon, head = _vl_canon(str(row.get("element", "")))
            if not canon or head == "O" or _is_ocr_artifact(head):  # O-баланс / OCR-мусор
                continue
            raw = str(row.get("value", "")).strip()
            rng = _cell_range(raw)
            if rng is not None:
                lo, hi = rng
            else:
                v = _value_num(raw)
                if v is None:
                    continue
                lo = hi = v
            comp = _cell_comparator(raw) or "="
            u = _vl_unit(row.get("unit"), head)
            if u == "pct" and hi is not None and hi > 100:   # санитарный гейт
                continue
            facts.append({
                "doc_id": doc_id, "node_type": "Material", "canon": canon,
                "metric": "содержание", "unit_canon": u,
                "value_low": lo, "value_high": hi, "comparator": comp,
                "confidence": 0.9, "phase": phase, "conditions": None,
                "source": "vision_vl",
                "quote": f"таблица состава{' [' + phase + ']' if phase else ''} | {head}: {raw}",
            })
    return facts


def extract_pdf_vl(path: str, doc_id: str, max_pages: int = 20):
    """VL-извлечение таблиц состава из PDF (для новых загрузок). Рендер fitz →
    пре-фильтр page_worth_ocr → VL → факты. VL выключен → []."""
    from src.config import VL_ENABLED
    if not VL_ENABLED:
        return []
    try:
        import fitz
    except Exception:  # noqa: BLE001
        return []
    facts = []
    try:
        d = fitz.open(path)
    except Exception:  # noqa: BLE001
        return []
    for pg_i in range(min(len(d), max_pages)):
        pg = d[pg_i]
        text = pg.get_text()
        if not page_worth_ocr(text):        # тот же пре-фильтр, что у OCR-пути
            continue
        png = pg.get_pixmap(dpi=180).tobytes("png")
        parsed = _vl_ocr(png)
        if parsed:
            facts += _vl_to_facts(parsed, doc_id, text[:200])
    return facts


# issue #4: пре-фильтр составности для DOCX-пути (у PDF роль играет page_worth_ocr).
_COMP_CTX_RE = re.compile(r"(содержани|масс|состав|wt|%|мас\.)", re.IGNORECASE)

def _table_is_composition(grid: dict, nrows: int, ncols: int) -> bool:
    """DOCX-таблица похожа на состав, если в её тексте есть маркеры состава
    ('содержание/масс/%/состав/wt') ЛИБО ≥50% числовых значений попадают в [0,100]."""
    all_text = " ".join(str(v) for v in grid.values())
    if _COMP_CTX_RE.search(all_text):
        return True
    nums = [n for n in (_num(v) for v in grid.values()) if n is not None]
    if not nums:
        return False
    in_range = sum(1 for n in nums if 0 <= n <= 100)
    return in_range >= len(nums) / 2

def extract_docx(path: str, doc_id: str):
    """DOCX-таблицы составов → факты структурно (без OCR) через ту же table_to_facts.
    python-docx: doc.tables[].rows[].cells[].text. Обзоры/Статьи в DOCX имеют
    таблицы состава, которые при плоском извлечении текста рассыпаются.
    Пре-фильтр составности (issue #4) отсекает нехимические таблицы (реестры/сметы)."""
    from docx import Document
    doc = Document(path)
    facts = []
    for tbl in doc.tables:
        rows = tbl.rows
        if not rows:
            continue
        ncols = max((len(r.cells) for r in rows), default=0)
        nrows = len(rows)
        if nrows < 2 or ncols < 2:
            continue
        grid = {}
        for ri, row in enumerate(rows):
            for ci, cell in enumerate(row.cells):
                grid[(ri, ci)] = (cell.text or "").replace("\n", " ").strip()
        if not _table_is_composition(grid, nrows, ncols):
            continue
        # контекст фазы — первая содержательная строка таблицы
        ctx = " ".join(grid.get((0, c), "") for c in range(ncols)).strip()[:80] or doc_id
        facts += table_to_facts(grid, nrows, ncols, doc_id, ctx)
    return facts

def main(limit: int = 0, min_kg: int = 3):
    if not LLM_ENABLED:
        print("Нет ключа — Vision OCR недоступен (ядро работает без него)"); return
    meta = [json.loads(l) for l in open(DOCS_META, encoding="utf-8")]
    from src.config import CORPUS_DIR
    # PDF (Vision OCR, пре-фильтр составных страниц) + DOCX (python-docx структурно)
    dense = [m for m in meta if (m.get("kg_value") or 0) >= min_kg and m.get("ok")
             and (m["src"].lower().endswith(".pdf") or m["src"].lower().endswith(".docx"))]
    if limit:
        dense = dense[:limit]
    npdf = sum(1 for m in dense if m["src"].lower().endswith(".pdf"))
    print(f"Доков для табличного извлечения: {len(dense)} (PDF {npdf}, DOCX {len(dense)-npdf})", flush=True)
    all_facts, t0 = [], time.time()
    for i, m in enumerate(dense, 1):
        src = m["src"].split(" :: ")[0]   # для файлов из архивов берём контейнер
        path = os.path.join(str(CORPUS_DIR), src)
        if not os.path.exists(path):
            continue
        try:
            if src.lower().endswith(".pdf"):
                all_facts += extract_pdf_vl(path, m["doc_id"])
            elif src.lower().endswith(".docx"):
                all_facts += extract_docx(path, m["doc_id"])
        except Exception as e:
            print(f"  err {m['src'][:40]}: {str(e)[:50]}")
        if i % 10 == 0:
            print(f"  {i}/{len(dense)} фактов={len(all_facts)} {time.time()-t0:.0f}s", flush=True)
    with open(VISION_FACTS, "w", encoding="utf-8") as fh:
        for f in all_facts:
            fh.write(json.dumps(f, ensure_ascii=False) + "\n")
    print(f"ГОТОВО: vision-фактов {len(all_facts)} → {VISION_FACTS}, {time.time()-t0:.0f}s")

if __name__ == "__main__":
    main(int(sys.argv[1]) if len(sys.argv) > 1 else 0)
